"""
自動化雲端出卷系統 (Automated Cloud Exam System)
使用 Streamlit + Google Sheets + Python

架構：
- 資料層：Google Sheets（題庫）
- 邏輯層：Streamlit（篩選、隨機、匯出）
- 部署層：Streamlit Cloud（自動部署）
"""

import streamlit as st
import pandas as pd
import random
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
import gspread
from oauth2client.service_account import ServiceAccountCredentials

# ==================== 頁面設定 ====================
st.set_page_config(
    page_title="自動化出卷系統",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 樣式設定 ====================
st.markdown("""
<style>
    .main-title {
        text-align: center;
        color: #1f77b4;
        margin-bottom: 20px;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        margin: 10px 0;
    }
    .question-card {
        background-color: #ffffff;
        border-left: 4px solid #1f77b4;
        padding: 15px;
        margin: 10px 0;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

# ==================== 快取設定 ====================
@st.cache_data(ttl=300)  # 5 分鐘快取
def load_google_sheets(sheet_id):
    """
    從 Google Sheets 讀取題庫
    支援公開試算表（無需認證）
    """
    try:
        # 使用 CSV 匯出方式讀取公開試算表
        url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
        df = pd.read_csv(url)
        
        # 資料驗證
        required_columns = ['ID', '類型', '科目', '題目內容', '參考解答', '分數']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            st.error(f"❌ 試算表缺少欄位：{', '.join(missing_columns)}")
            st.info("📋 必要欄位：ID、類型、科目、題目內容、參考解答、分數")
            return None
        
        # 資料清理
        df['分數'] = pd.to_numeric(df['分數'], errors='coerce').fillna(25).astype(int)
        df = df.dropna(subset=['題目內容'])
        
        return df
    
    except Exception as e:
        st.error(f"❌ 連接失敗：{str(e)}")
        st.info("💡 提示：確認試算表已分享為「任何有連結的人都可以檢視」")
        return None

# ==================== 核心邏輯 ====================
def generate_exam(df, target_score):
    """
    隨機生成符合目標分數的考卷
    
    演算法：
    1. 隨機排序題目
    2. 貪心選擇：逐題加入，直到達到目標分數
    3. 返回選中題目和實際分數
    """
    if df.empty:
        return None
    
    selected_questions = []
    current_score = 0
    
    # 隨機排序
    shuffled_df = df.sample(frac=1).reset_index(drop=True)
    
    # 貪心選擇
    for _, row in shuffled_df.iterrows():
        question_score = int(row['分數'])
        if current_score + question_score <= target_score:
            selected_questions.append(row)
            current_score += question_score
    
    if not selected_questions:
        return None
    
    return {
        'questions': selected_questions,
        'total_score': current_score,
        'target_score': target_score,
        'question_count': len(selected_questions),
        'generated_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

# ==================== 匯出功能 ====================
def export_to_word(exam_data):
    """匯出為 Word 檔 (.docx)"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('法律考試考卷', 0)
    title_format = title.paragraph_format
    title_format.alignment = 1  # 置中
    
    # 基本資訊
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.cell(0, 0).text = '生成時間'
    info_table.cell(0, 1).text = exam_data['generated_time']
    info_table.cell(1, 0).text = '總分'
    info_table.cell(1, 1).text = f"{exam_data['total_score']} 分"
    info_table.cell(2, 0).text = '題數'
    info_table.cell(2, 1).text = f"{exam_data['question_count']} 題"
    
    doc.add_paragraph()
    
    # 題目
    for idx, question in enumerate(exam_data['questions'], 1):
        # 題號和分數
        heading = doc.add_heading(f"第 {idx} 題 ({int(question['分數'])} 分)", level=2)
        
        # 題目資訊
        info_para = doc.add_paragraph()
        info_para.add_run('科目：').bold = True
        info_para.add_run(f"{question['科目']} | ")
        info_para.add_run('類型：').bold = True
        info_para.add_run(f"{question['類型']}")
        
        # 題目內容
        content_para = doc.add_paragraph()
        content_para.add_run('題目：').bold = True
        content_para.add_run(f"\n{question['題目內容']}")
        
        # 參考解答
        if pd.notna(question['參考解答']) and str(question['參考解答']).strip():
            answer_para = doc.add_paragraph()
            answer_para.add_run('參考解答：').bold = True
            answer_para.add_run(f"\n{question['參考解答']}")
            answer_para.paragraph_format.left_indent = Inches(0.5)
        
        # 答題空間
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
    
    # 保存到記憶體
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def export_to_pdf(exam_data):
    """匯出為 PDF 檔"""
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # 自訂樣式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f77b4'),
        spaceAfter=30,
        alignment=1,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1f77b4'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    story = []
    
    # 標題
    story.append(Paragraph("法律考試考卷", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # 基本資訊表格
    info_data = [
        ['生成時間', exam_data['generated_time']],
        ['總分', f"{exam_data['total_score']} 分"],
        ['題數', f"{exam_data['question_count']} 題"]
    ]
    
    info_table = Table(info_data, colWidths=[1.5*inch, 3.5*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8f0f8')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 0.3*inch))
    
    # 題目
    for idx, question in enumerate(exam_data['questions'], 1):
        # 題號
        story.append(Paragraph(
            f"第 {idx} 題 ({int(question['分數'])} 分)",
            heading_style
        ))
        
        # 科目和類型
        story.append(Paragraph(
            f"<b>科目：</b>{question['科目']} | <b>類型：</b>{question['類型']}",
            styles['Normal']
        ))
        
        # 題目內容
        story.append(Paragraph(
            f"<b>題目：</b>",
            styles['Normal']
        ))
        story.append(Paragraph(
            str(question['題目內容']),
            styles['Normal']
        ))
        
        # 參考解答
        if pd.notna(question['參考解答']) and str(question['參考解答']).strip():
            story.append(Paragraph(
                f"<b>參考解答：</b>",
                styles['Normal']
            ))
            story.append(Paragraph(
                str(question['參考解答']),
                styles['Normal']
            ))
        
        story.append(Spacer(1, 0.2*inch))
        
        # 分頁
        if idx < len(exam_data['questions']):
            story.append(PageBreak())
    
    doc.build(story)
    output.seek(0)
    return output.getvalue()

# ==================== UI 元件 ====================
def display_exam(exam_data):
    """顯示生成的考卷"""
    # 考卷頭部
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("總分", f"{exam_data['total_score']} 分")
    with col2:
        st.metric("題數", f"{exam_data['question_count']} 題")
    with col3:
        st.metric("生成時間", exam_data['generated_time'].split()[0])
    
    st.divider()
    
    # 題目顯示
    for idx, question in enumerate(exam_data['questions'], 1):
        with st.container(border=True):
            # 題號和分數
            st.markdown(f"### 第 {idx} 題 ({int(question['分數'])} 分)")
            
            # 科目和類型
            col1, col2 = st.columns(2)
            with col1:
                st.caption(f"📚 科目：{question['科目']}")
            with col2:
                st.caption(f"📝 類型：{question['類型']}")
            
            # 題目內容
            st.markdown("**題目內容：**")
            st.markdown(f"> {question['題目內容']}")
            
            # 參考解答（可展開）
            if pd.notna(question['參考解答']) and str(question['參考解答']).strip():
                with st.expander("📖 查看參考解答"):
                    st.markdown(f"{question['參考解答']}")
    
    st.divider()
    
    # 匯出功能
    st.subheader("📥 匯出考卷")
    col1, col2 = st.columns(2)
    
    with col1:
        word_data = export_to_word(exam_data)
        st.download_button(
            label="📄 下載為 Word (.docx)",
            data=word_data,
            file_name=f"考卷_{exam_data['generated_time'].replace(':', '-').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col2:
        pdf_data = export_to_pdf(exam_data)
        st.download_button(
            label="📋 下載為 PDF",
            data=pdf_data,
            file_name=f"考卷_{exam_data['generated_time'].replace(':', '-').replace(' ', '_')}.pdf",
            mime="application/pdf",
            use_container_width=True
        )

# ==================== 主程式 ====================
def main():
    # 頁面標題
    st.markdown("<h1 class='main-title'>📝 自動化雲端出卷系統</h1>", unsafe_allow_html=True)
    st.markdown("基於 Google Sheets 題庫的智能出卷平台 | 資料與程式完全分離", unsafe_allow_html=True)
    st.divider()
    
    # 側邊欄設定
    with st.sidebar:
        st.header("⚙️ 系統設定")
        
        # Google Sheets ID 輸入
        sheet_id = st.text_input(
            "📊 Google Sheets ID",
            placeholder="貼上你的試算表 ID",
            help="從 Google Sheets 分享連結中複製 ID\nhttps://docs.google.com/spreadsheets/d/[ID]/edit"
        )
        
        if not sheet_id:
            st.info("💡 如何取得 Google Sheets ID？\n\n1. 打開你的 Google Sheets\n2. 點擊「分享」\n3. 確保設為「任何有連結的人都可以檢視」\n4. 從連結中複製 ID")
            return
        
        # 讀取題庫
        df = load_google_sheets(sheet_id)
        
        if df is None or df.empty:
            st.error("❌ 無法讀取題庫，請檢查 ID 和分享設定")
            return
        
        st.success(f"✅ 成功讀取 {len(df)} 題")
        
        # 題庫統計
        st.subheader("📊 題庫統計")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("總題數", len(df))
        with col2:
            st.metric("總分數", int(df['分數'].sum()))
        
        # 篩選條件
        st.subheader("🔍 篩選條件")
        
        # 科目篩選
        subjects = sorted(df['科目'].unique().tolist())
        selected_subjects = st.multiselect(
            "選擇科目",
            subjects,
            default=subjects,
            help="選擇要包含在考卷中的科目"
        )
        
        # 題型篩選
        types = sorted(df['類型'].unique().tolist())
        selected_types = st.multiselect(
            "選擇題型",
            types,
            default=types,
            help="選擇要包含在考卷中的題型"
        )
        
        # 分數篩選
        min_score = int(df['分數'].min())
        max_score = int(df['分數'].max())
        score_range = st.slider(
            "分數範圍",
            min_value=min_score,
            max_value=max_score,
            value=(min_score, max_score),
            help="選擇題目分數範圍"
        )
        
        # 目標總分
        target_score = st.number_input(
            "🎯 目標總分",
            min_value=25,
            max_value=500,
            value=100,
            step=25,
            help="生成的考卷會盡量接近此分數"
        )
        
        # 應用篩選
        filtered_df = df[
            (df['科目'].isin(selected_subjects)) &
            (df['類型'].isin(selected_types)) &
            (df['分數'] >= score_range[0]) &
            (df['分數'] <= score_range[1])
        ]
        
        # 篩選結果
        st.info(f"符合條件的題目：**{len(filtered_df)}** 題 / 總分 **{int(filtered_df['分數'].sum())}** 分")
        
        # 生成考卷按鈕
        if st.button("🎲 隨機生成考卷", use_container_width=True, type="primary"):
            if len(filtered_df) == 0:
                st.error("❌ 沒有符合條件的題目，請調整篩選條件")
            else:
                exam = generate_exam(filtered_df, target_score)
                if exam is None:
                    st.error("❌ 無法生成考卷，請調整目標分數或篩選條件")
                else:
                    st.session_state.generated_exam = exam
                    st.session_state.show_exam = True
                    st.rerun()
    
    # 主要內容區
    if 'show_exam' in st.session_state and st.session_state.show_exam:
        display_exam(st.session_state.generated_exam)
    else:
        # 歡迎頁面
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("""
            ## 👋 歡迎使用自動化出卷系統
            
            這是一個基於 **Google Sheets** 的智能出卷平台。
            
            ### 🎯 核心功能
            
            - **📊 題庫管理**：在 Google Sheets 中管理所有考題
            - **🔍 智能篩選**：按科目、題型、分數篩選題目
            - **🎲 隨機出卷**：自動生成符合目標分數的考卷
            - **📥 多格式匯出**：支援 Word 和 PDF 格式
            
            ### 🚀 快速開始
            
            1. **準備 Google Sheets**
               - 建立試算表，包含欄位：ID、類型、科目、題目內容、參考解答、分數
               - 分享為「任何有連結的人都可以檢視」
            
            2. **輸入試算表 ID**
               - 在左側邊欄輸入你的 Google Sheets ID
            
            3. **設定篩選條件**
               - 選擇科目、題型、分數範圍
               - 設定目標總分
            
            4. **生成考卷**
               - 點擊「隨機生成考卷」按鈕
               - 下載為 Word 或 PDF
            
            ### 💡 小提示
            
            - 題庫更新後，重新整理頁面即可看到新題目
            - 程式碼更新後，自動部署到 Streamlit Cloud
            - 題庫和程式完全分離，互不影響
            """)
        
        with col2:
            st.markdown("""
            ### 📚 支援的科目
            
            - 民法
            - 民事訴訟法
            - 刑法
            - 刑事訴訟法
            - 行政法
            - 憲法
            - 商事法
            - 其他...
            
            ### 📝 支援的題型
            
            - 申論題
            - 實例題
            - 選擇題
            - 其他...
            """)

# ==================== 進入點 ====================
if __name__ == "__main__":
    # 初始化 session state
    if 'show_exam' not in st.session_state:
        st.session_state.show_exam = False
    if 'generated_exam' not in st.session_state:
        st.session_state.generated_exam = None
    
    main()
